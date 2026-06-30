# -*- coding: utf-8 -*-
"""
사용법_Avikus_Newsletter.md  ->  사용법_Avikus_Newsletter.docx 변환기.
markdown 의 일부 문법(제목/목록/표/코드블록/굵게/인라인코드/구분선)을 Word 로 변환합니다.
실행:  py build_docx.py
필요:  py -m pip install python-docx
"""
import os, re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
MD = os.path.join(BASE, "사용법_Avikus_Newsletter.md")
OUT = os.path.join(BASE, "사용법_Avikus_Newsletter.docx")

BODY_FONT = "맑은 고딕"     # 한글 본문 폰트
CODE_FONT = "Consolas"
NAVY = RGBColor(0x1F, 0x3A, 0x5F)

def set_font(run, name=BODY_FONT, size=None, bold=None, color=None, mono=False):
    f = run.font
    f.name = CODE_FONT if mono else name
    # 한글 폰트도 같이 지정 (East Asian)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None: f.size = Pt(size)
    if bold is not None: f.bold = bold
    if color is not None: f.color.rgb = color

def add_inline(p, text):
    """**굵게** 와 `인라인코드` 를 파싱해 run 으로 추가."""
    # 토큰 분리: **...**  또는  `...`
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2]); set_font(r, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1]); set_font(r, mono=True, color=NAVY)
        else:
            r = p.add_run(part); set_font(r)

def shade(cell_or_par_elm, hex_fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_fill)
    cell_or_par_elm.append(shd)

def add_code_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.1)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(8)
    shade(p._p.get_or_add_pPr(), "F2F4F7")
    for i, ln in enumerate(lines):
        if i > 0:
            r = p.add_run(); r.add_break()
        r = p.add_run(ln); set_font(r, mono=True, size=9.5)

def add_table(doc, rows):
    # rows: list of list[str]; 첫 행은 헤더
    ncol = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=ncol)
    t.style = "Light Grid Accent 1"
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci in range(ncol):
            txt = row[ci] if ci < len(row) else ""
            cell = cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            add_inline(p, txt)
            if ri == 0:
                for run in p.runs:
                    run.font.bold = True
    return t

def split_row(line):
    line = line.strip()
    if line.startswith("|"): line = line[1:]
    if line.endswith("|"): line = line[:-1]
    return [c.strip() for c in line.split("|")]

def main():
    with open(MD, encoding="utf-8") as f:
        lines = f.read().split("\n")

    doc = Document()
    # 기본 스타일 폰트
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT; normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # 코드블록
        if stripped.startswith("```"):
            j = i + 1; buf = []
            while j < n and not lines[j].strip().startswith("```"):
                buf.append(lines[j]); j += 1
            add_code_block(doc, buf)
            i = j + 1; continue

        # 표 (다음 줄이 |---| 구분선)
        if stripped.startswith("|") and i + 1 < n and re.match(r"^\|?\s*:?-{2,}", lines[i+1].strip()):
            rows = [split_row(stripped)]
            j = i + 2
            while j < n and lines[j].strip().startswith("|"):
                rows.append(split_row(lines[j])); j += 1
            add_table(doc, rows)
            doc.add_paragraph()
            i = j; continue

        # 빈 줄
        if stripped == "":
            i += 1; continue

        # 구분선
        if stripped == "---":
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
            pPr = p._p.get_or_add_pPr()
            pbdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "1F3A5F")
            pbdr.append(bottom); pPr.append(pbdr)
            i += 1; continue

        # 제목
        if stripped.startswith("### "):
            h = doc.add_heading(level=2); add_inline(h, stripped[4:])
            i += 1; continue
        if stripped.startswith("## "):
            h = doc.add_heading(level=1); add_inline(h, stripped[3:])
            i += 1; continue
        if stripped.startswith("# "):
            h = doc.add_heading(level=0); add_inline(h, stripped[2:])
            i += 1; continue

        # 인용
        if stripped.startswith("> "):
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.3)
            r = p.add_run("│ "); set_font(r, color=NAVY, bold=True)
            add_inline(p, stripped[2:])
            i += 1; continue

        # 순서 목록
        m = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if m:
            p = doc.add_paragraph(style="List Number"); add_inline(p, m.group(2))
            i += 1; continue

        # 글머리 목록
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet"); add_inline(p, stripped[2:])
            i += 1; continue

        # 일반 문단 (이탤릭 *...* 한 줄 처리 포함)
        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            p = doc.add_paragraph(); r = p.add_run(stripped[1:-1]); set_font(r); r.font.italic = True
            i += 1; continue

        p = doc.add_paragraph(); add_inline(p, stripped)
        i += 1

    doc.save(OUT)
    print("생성 완료:", OUT)

if __name__ == "__main__":
    main()
